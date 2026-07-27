import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- EXACT COLOR PALETTE FROM CLAUDE'S SCRIPT ---
BLUE = RGBColor(0x1F, 0x4E, 0x79)       # "1F4E79"
LIGHT_BLUE = RGBColor(0x2E, 0x75, 0xB6) # "2E75B6"
GRAY = RGBColor(0x59, 0x59, 0x59)       # "595959"
LIGHT_GRAY = RGBColor(0x77, 0x77, 0x77) # "777777"
LINE_COLOR_HEX = "2E75B6"

def set_paragraph_bottom_border(paragraph, color_hex=LINE_COLOR_HEX, size="6"):
    """Adds a bottom border to a paragraph (equivalent to BorderStyle.SINGLE in docx.js)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def apply_default_margins(doc):
    """Apply uniform margins to the entire document."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

# --- COMPONENTES ATÓMICOS REUTILIZABLES ---

def add_main_title(doc, text: str, subtitle: str = "", contact_or_meta: str = ""):
    """Standard main header for any document."""
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run(text.upper())
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = BLUE

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.name = "Arial"
        r_sub.font.size = Pt(11)
        r_sub.font.color.rgb = LIGHT_BLUE

    if contact_or_meta:
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_meta.paragraph_format.space_after = Pt(12)
        r_meta = p_meta.add_run(contact_or_meta)
        r_meta.font.name = "Arial"
        r_meta.font.size = Pt(9)
        r_meta.font.color.rgb = GRAY

def section_header(doc, text: str):
    """Generates a section header with an underline border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14) # ~280 dxa
    p.paragraph_format.space_after = Pt(4)   # ~80 dxa
    set_paragraph_bottom_border(p, color_hex=LINE_COLOR_HEX, size="6")
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13) # ~26 pt in docx.js
    run.font.color.rgb = BLUE
    return p

def add_paragraph_body(doc, text: str, bold_prefix: str = ""):
    """Adds a standardized body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.color.rgb = BLUE
        
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def bullet(doc, text: str):
    """Adds a bullet list item with proper left indentation."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2) # ~40 dxa
    p.paragraph_format.space_after = Pt(2)  # ~40 dxa
    
    # --- INDENTATION ADJUSTMENT ---
    # left_indent: Espacio total desde el margen izquierdo hasta el texto (~0.25 in / ~360 dxa)
    # first_line_indent: Posición del bullet respecto al texto (-0.15 in / ~216 dxa)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10) # ~20 pt
    run.font.color.rgb = GRAY
    return p

def job_title(doc, company: str, role: str, dates: str):
    """Generates a job title row with company, role, and right-aligned dates."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9) # ~180 dxa
    p.paragraph_format.space_after = Pt(2)  # ~40 dxa
    
    # Set a right-aligned tab stop at the right margin (~6.27 inches / 9026 dxa)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.27), WD_TAB_ALIGNMENT.RIGHT)
    
    r_comp = p.add_run(company)
    r_comp.bold = True
    r_comp.font.name = "Arial"
    r_comp.font.size = Pt(11) # ~22 pt
    r_comp.font.color.rgb = BLUE

    r_sep = p.add_run(" | ")
    r_sep.font.name = "Arial"
    r_sep.font.size = Pt(11)
    r_sep.font.color.rgb = GRAY

    r_role = p.add_run(role)
    r_role.font.name = "Arial"
    r_role.font.size = Pt(11)
    r_role.font.color.rgb = GRAY

    r_tab = p.add_run("\t")
    r_tab.font.name = "Arial"
    r_tab.font.size = Pt(10)

    r_dates = p.add_run(dates)
    r_dates.italic = True
    r_dates.font.name = "Arial"
    r_dates.font.size = Pt(10)
    r_dates.font.color.rgb = GRAY
    return p

def competencia(doc, title: str, description: str):
    """Generates a competency block with title and indented description."""
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6) # ~120 dxa
    p_title.paragraph_format.space_after = Pt(1)  # ~20 dxa
    
    r_t = p_title.add_run(title)
    r_t.bold = True
    r_t.font.name = "Arial"
    r_t.font.size = Pt(10)
    r_t.font.color.rgb = LIGHT_BLUE

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after = Pt(4) # ~80 dxa
    p_desc.paragraph_format.left_indent = Inches(0.25) # ~360 dxa
    
    r_d = p_desc.add_run(description)
    r_d.font.name = "Arial"
    r_d.font.size = Pt(10)
    r_d.font.color.rgb = GRAY

# --- TEMPLATE 1: GENERIC DOCUMENT / REPORT (Default) ---
def generate_default_document(
    output_filename: str,
    title: str,
    subtitle: str = "",
    meta_info: str = "",
    sections: list[dict] = None
) -> str:
    """
    Generates a generic document/report (technical report, summary, note) with the corporate design.
    
    Args:
        output_filename: Name of the output file.
        title: Main title of the document.
        subtitle: Optional subtitle.
        meta_info: Header information (date, author, version).
        sections: List of sections with 'title', 'paragraphs' (list of text) and 'bullets' (list of points).
    """
    if not output_filename.endswith('.docx'):
        output_filename = f"{output_filename.rsplit('.', 1)[0]}.docx"

    doc = Document()
    apply_default_margins(doc)

    # Cabecera
    add_main_title(doc, title, subtitle, meta_info)

    # Contenido dinámico
    if sections:
        for sec in sections:
            sec_title = sec.get("title", "")
            if sec_title:
                section_header(doc, sec_title)
            
            for p_text in sec.get("paragraphs", []):
                add_paragraph_body(doc, p_text)
                
            for b_text in sec.get("bullets", []):
                bullet(doc, b_text)

    doc.save(output_filename)
    final_path = convert_docx_to_pdf(output_filename)
    return f"Generic document created successfully at: {os.path.abspath(final_path)}"

def generate_cv_document(
    output_filename: str,
    full_name: str,
    sub_title: str,
    contact_info: str,
    profile_summary: str,
    competencies: list[dict],
    experience: list[dict],
    education: list[str],
    languages: str
) -> str:
    """
    Generates a styled CV document (.docx and converted .pdf) following the exact corporate template.
    
    Args:
        output_filename: Name of the file to save (e.g. CV_Raul_Marchan.docx)
        full_name: Candidate full name for the header
        sub_title: Professional title / Role
        contact_info: Contact line separated by pipes
        profile_summary: Paragraph summarizing the candidate profile
        competencies: List of dicts with 'title' and 'description'
        experience: List of dicts with 'company', 'role', 'dates', and 'bullets' (list of strings)
        education: List of strings for education/certifications
        languages: Languages summary string
    """
    if not output_filename.endswith('.docx'):
        output_filename = f"{output_filename.rsplit('.', 1)[0]}.docx"

    doc = Document()

    # Page Margins (0.5 in top/bottom, 0.6 in left/right)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # 1. HEADER
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(full_name.upper())
    r_name.bold = True
    r_name.font.name = "Arial"
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)
    r_sub = p_sub.add_run(sub_title)
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = LIGHT_BLUE

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    r_contact = p_contact.add_run(contact_info)
    r_contact.font.name = "Arial"
    r_contact.font.size = Pt(9)
    r_contact.font.color.rgb = GRAY

    # 2. PROFESSIONAL PROFILE
    if profile_summary:
        section_header(doc, "PROFESSIONAL PROFILE")
        p_prof = doc.add_paragraph()
        p_prof.paragraph_format.space_before = Pt(4)
        p_prof.paragraph_format.space_after = Pt(6)
        r_prof = p_prof.add_run(profile_summary)
        r_prof.font.name = "Arial"
        r_prof.font.size = Pt(10)
        r_prof.font.color.rgb = GRAY

    # 3. ACCREDITED COMPETENCIES
    if competencies:
        section_header(doc, "ACCREDITED PROFESSIONAL COMPETENCIES")
        for comp in competencies:
            competencia(doc, comp.get("title", ""), comp.get("description", ""))

    # 4. PROFESSIONAL EXPERIENCE
    if experience:
        section_header(doc, "PROFESSIONAL EXPERIENCE")
        for job in experience:
            job_title(doc, job.get("company", ""), job.get("role", ""), job.get("dates", ""))
            for b in job.get("bullets", []):
                bullet(doc, b)

    # 5. EDUCATION & CERTIFICATIONS
    if education:
        section_header(doc, "EDUCATION & CERTIFICATIONS")
        for edu_item in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(3)
            p_edu.paragraph_format.space_after = Pt(3)
            r_edu = p_edu.add_run(edu_item)
            r_edu.font.name = "Arial"
            r_edu.font.size = Pt(10)
            r_edu.font.color.rgb = GRAY

    # 6. LANGUAGES
    if languages:
        section_header(doc, "LANGUAGES")
        p_lang = doc.add_paragraph()
        p_lang.paragraph_format.space_before = Pt(4)
        p_lang.paragraph_format.space_after = Pt(4)
        r_lang = p_lang.add_run(languages)
        r_lang.font.name = "Arial"
        r_lang.font.size = Pt(10)
        r_lang.font.color.rgb = GRAY

    # Save document
    doc.save(output_filename)

    # Auto-convert to PDF
    final_path = convert_docx_to_pdf(output_filename)
    return f"Document successfully generated at: {os.path.abspath(final_path)}"


def convert_docx_to_pdf(docx_path: str) -> str:
    """Converts a .docx file to .pdf using available local engines."""
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    # Attempt 1: docx2pdf (macOS/Windows Word Automation)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # Attempt 2: Pandoc CLI
    try:
        subprocess.run(['pandoc', docx_path, '-o', pdf_path], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    print("ℹ️ Generated .docx file. (Install docx2pdf or pandoc to enable automatic PDF export).")
    return docx_path